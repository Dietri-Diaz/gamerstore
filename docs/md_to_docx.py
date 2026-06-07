# Conversor simple Markdown -> Word (.docx) para la documentacion APF3.
# Maneja: titulos (#..####), parrafos, listas (- y 1.), tablas (| |) y bloques de codigo (```).
import re
from docx import Document
from docx.shared import Pt

SRC = r"C:\Users\dietr\Desktop\Pruebas\UNIVERSIDAD\gamerstore\docs\APF3-Documentacion-GamerStore.md"
OUT = r"C:\Users\dietr\Desktop\Pruebas\UNIVERSIDAD\gamerstore\docs\APF3-Documentacion-GamerStore.docx"

doc = Document()
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

INLINE = re.compile(r'(\*\*.+?\*\*|`.+?`)')

def add_runs(p, text):
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(9.5)
        else:
            p.add_run(part)

def clean(text):
    return text.replace('**', '').replace('`', '').replace('\\|', '|')

def is_sep(s):
    return bool(re.match(r'^\s*\|?[\s:\|-]+\|?\s*$', s)) and '-' in s

def split_row(s):
    s = s.strip().strip('|').replace('\\|', '\x00')
    return [c.strip().replace('\x00', '|') for c in s.split('|')]

lines = open(SRC, encoding='utf-8').read().split('\n')
i = 0
while i < len(lines):
    line = lines[i]

    if line.strip() == '---':
        i += 1; continue

    if line.strip().startswith('```'):
        i += 1
        code = []
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code.append(lines[i]); i += 1
        i += 1
        p = doc.add_paragraph()
        run = p.add_run('\n'.join(code))
        run.font.name = 'Consolas'; run.font.size = Pt(9)
        continue

    m = re.match(r'^(#{1,6})\s+(.*)', line)
    if m:
        doc.add_heading(clean(m.group(2)), level=min(len(m.group(1)), 4))
        i += 1; continue

    if line.strip().startswith('|') and i + 1 < len(lines) and is_sep(lines[i+1]):
        header = split_row(line)
        i += 2
        rows = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            rows.append(split_row(lines[i])); i += 1
        t = doc.add_table(rows=1, cols=len(header))
        t.style = 'Table Grid'
        for j, h in enumerate(header):
            t.rows[0].cells[j].text = clean(h)
        for r in rows:
            cells = t.add_row().cells
            for j in range(len(header)):
                cells[j].text = clean(r[j]) if j < len(r) else ''
        doc.add_paragraph()
        continue

    m = re.match(r'^\s*-\s+(.*)', line)
    if m:
        add_runs(doc.add_paragraph(style='List Bullet'), m.group(1))
        i += 1; continue

    m = re.match(r'^\s*\d+\.\s+(.*)', line)
    if m:
        add_runs(doc.add_paragraph(style='List Number'), m.group(1))
        i += 1; continue

    if line.strip() == '':
        i += 1; continue

    add_runs(doc.add_paragraph(), line)
    i += 1

doc.save(OUT)
print("OK ->", OUT)
