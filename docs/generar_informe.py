# -*- coding: utf-8 -*-
# Convierte INFORME-PROYECTO.md en un Word (.docx) formateado con portada, indice y estilos.
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = r"C:\Users\dietr\Desktop\Pruebas\UNIVERSIDAD\gamerstore-main\docs\INFORME-PROYECTO.md"
OUT = r"C:\Users\dietr\Desktop\Pruebas\UNIVERSIDAD\gamerstore-main\docs\GamerStore-Informe-Proyecto.docx"

AZUL = RGBColor(0x1F, 0x38, 0x64)
ACENTO = RGBColor(0x4F, 0x46, 0xE5)

doc = Document()
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

# ---------- PORTADA ----------
def linea(txt, size, bold=False, color=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(txt)
    r.font.size = Pt(size); r.bold = bold
    if color is not None:
        r.font.color.rgb = color
    return p

linea("Universidad Tecnológica del Perú (UTP)", 13, bold=True, color=AZUL, space_before=60)
linea("Facultad de Ingeniería", 11)
doc.add_paragraph()
linea("GamerStore", 40, bold=True, color=ACENTO, space_before=40)
linea("Sistema de Administración (ERP) para una Tienda de Tecnología", 15, bold=True, color=AZUL)
doc.add_paragraph()
linea("INFORME DEL PROYECTO", 16, bold=True, space_before=20)
doc.add_paragraph()
linea("Spring Boot (API REST)  ·  React (Vite)  ·  MySQL  ·  Seguridad con JWT", 11, color=RGBColor(0x64,0x74,0x8B))
for _ in range(6):
    doc.add_paragraph()
for etiqueta in ["Curso:  ______________________________",
                 "Docente:  ____________________________",
                 "Integrantes:  ________________________",
                 "                     ________________________",
                 "                     ________________________",
                 "Fecha:  ______________________________"]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(etiqueta); r.font.size = Pt(11)
doc.add_page_break()

# ---------- CONTENIDO ----------
h = doc.add_heading('Contenido', level=1)
for run in h.runs: run.font.color.rgb = AZUL
secciones = [
    "1. Introducción", "2. Objetivos", "3. Alcance del sistema", "4. Tecnologías utilizadas",
    "5. Arquitectura del sistema", "6. Modelo de datos", "7. Seguridad",
    "8. Módulos del sistema (ERP)", "9. Integraciones externas", "10. Flujo general de una petición",
    "11. Datos de prueba (carga inicial)", "12. Pruebas y verificación", "13. Instalación y ejecución",
    "14. Conclusiones", "15. Anexos",
]
for s in secciones:
    doc.add_paragraph(s, style='List Bullet')
doc.add_page_break()

# ---------- CUERPO (desde el markdown) ----------
INLINE = re.compile(r'(\*\*.+?\*\*|`.+?`)')

def add_runs(p, text):
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(9.5)
        else:
            p.add_run(part)

def clean(text):
    return text.replace('**', '').replace('`', '').replace('\\|', '|')

def is_sep(s):
    return bool(re.match(r'^\s*\|?[\s:\|-]+\|?\s*$', s)) and '-' in s

def split_row(s):
    s = s.strip().strip('|').replace('\\|', '\x00')
    return [c.strip().replace('\x00', '|') for c in s.split('|')]

lines = open(SRC, encoding='utf-8').read().split('\n')
i = 0
while i < len(lines):
    line = lines[i]

    if line.strip() == '---':
        i += 1; continue

    if line.strip().startswith('```'):
        i += 1; code = []
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code.append(lines[i]); i += 1
        i += 1
        p = doc.add_paragraph()
        r = p.add_run('\n'.join(code)); r.font.name = 'Consolas'; r.font.size = Pt(9)
        continue

    m = re.match(r'^(#{1,6})\s+(.*)', line)
    if m:
        nivel = max(1, len(m.group(1)) - 1)  # '##' -> Heading 1
        head = doc.add_heading(clean(m.group(2)), level=min(nivel, 4))
        for run in head.runs: run.font.color.rgb = AZUL
        i += 1; continue

    if line.strip().startswith('|') and i + 1 < len(lines) and is_sep(lines[i+1]):
        header = split_row(line); i += 2; rows = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            rows.append(split_row(lines[i])); i += 1
        t = doc.add_table(rows=1, cols=len(header)); t.style = 'Light Grid Accent 1'
        for j, hh in enumerate(header):
            cell = t.rows[0].cells[j]; cell.text = clean(hh)
            for pr in cell.paragraphs:
                for run in pr.runs: run.bold = True
        for row in rows:
            cells = t.add_row().cells
            for j in range(len(header)):
                cells[j].text = clean(row[j]) if j < len(row) else ''
        doc.add_paragraph()
        continue

    m = re.match(r'^\s*-\s+(.*)', line)
    if m:
        add_runs(doc.add_paragraph(style='List Bullet'), m.group(1)); i += 1; continue

    m = re.match(r'^\s*\d+\.\s+(.*)', line)
    if m:
        add_runs(doc.add_paragraph(style='List Number'), m.group(1)); i += 1; continue

    if line.strip() == '':
        i += 1; continue

    add_runs(doc.add_paragraph(), line); i += 1

doc.save(OUT)
print("OK ->", OUT)
